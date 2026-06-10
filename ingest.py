"""
文档导入 + 向量化
把 documents/ 目录下的文档向量化并存入 ChromaDB
"""

import os
import shutil
from pathlib import Path

from langchain_community.document_loaders import (
    PyPDFLoader,
    TextLoader,
    UnstructuredMarkdownLoader,
)
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_openai import OpenAIEmbeddings
from langchain_chroma import Chroma

from config import (
    DOCUMENTS_DIR,
    VECTOR_STORE_DIR,
    EMBEDDING_PROVIDER,
    OPENAI_API_KEY,
    OPENAI_EMBEDDING_MODEL,
    LOCAL_EMBEDDING_MODEL,
    CHUNK_SIZE,
    CHUNK_OVERLAP,
)

# 全局 embedding 缓存（与 query.py 共享）
_embedding_model_cache = None


def get_embedding_model():
    """获取 Embedding 模型（带缓存）"""
    global _embedding_model_cache
    if _embedding_model_cache is not None:
        return _embedding_model_cache

    if EMBEDDING_PROVIDER == "openai":
        if not OPENAI_API_KEY:
            raise ValueError("🔴 使用 OpenAI Embedding 但未设置 OPENAI_API_KEY")
        _embedding_model_cache = OpenAIEmbeddings(
            model=OPENAI_EMBEDDING_MODEL,
            openai_api_key=OPENAI_API_KEY,
        )
    else:
        os.environ.setdefault("TRANSFORMERS_VERBOSITY", "error")
        _embedding_model_cache = HuggingFaceEmbeddings(
            model_name=LOCAL_EMBEDDING_MODEL,
            model_kwargs={
                "device": "cpu",
                "local_files_only": True,
            },
            encode_kwargs={"normalize_embeddings": True},
        )
    return _embedding_model_cache


def clear_embedding_cache():
    """清空 embedding 缓存（文档更新时需要）"""
    global _embedding_model_cache
    _embedding_model_cache = None


def _load_docx(file_path: str):
    """使用 python-docx 加载 DOCX 文件（比 Unstructured 更稳更快）"""
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        # 提取所有段落文字
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # 也提取表格中的文字
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text for cell in row.cells if cell.text.strip())
                if row_text.strip():
                    paragraphs.append(row_text)
        full_text = "\n".join(paragraphs) if paragraphs else ""
        if not full_text.strip():
            return []
        return [Document(page_content=full_text, metadata={"source": os.path.basename(file_path)})]
    except Exception as e:
        print(f"    ❌ DOCX 加载失败: {e}")
        return []


def load_documents(docs_dir: str):
    """加载目录下所有支持的文档（自动去重）"""
    docs_dir = Path(docs_dir)
    if not docs_dir.exists():
        print(f"创建文档目录: {docs_dir}")
        docs_dir.mkdir(parents=True, exist_ok=True)
        return []

    all_documents = []
    seen_paths = set()  # 用绝对路径去重（解决 Windows 大小写不敏感问题）

    supported_extensions = {
        ".pdf": PyPDFLoader,
        ".txt": TextLoader,
        ".md": UnstructuredMarkdownLoader,
        ".docx": "docx",  # 使用自定义 loader
    }

    for ext, loader_class in supported_extensions.items():
        # Windows 上 glob 可能大小写不敏感，只用小写搜一次
        for file_path in docs_dir.glob(f"*{ext}"):
            abs_path = str(file_path.resolve())
            if abs_path in seen_paths:
                continue
            seen_paths.add(abs_path)

            try:
                print(f"  加载: {file_path.name}")

                # 自定义 DOCX 加载器（使用 python-docx，比 Unstructured 快且稳定）
                if loader_class == "docx":
                    documents = _load_docx(str(file_path))
                else:
                    loader = loader_class(str(file_path))
                    documents = loader.load()

                for doc in documents:
                    doc.metadata["source"] = file_path.name
                all_documents.extend(documents)
                print(f"    -> {len(documents)} 页/段")
            except Exception as e:
                print(f"    ❌ 失败: {e}")

    if not all_documents:
        print("  (没有找到支持的文档)")

    return all_documents


def split_documents(documents):
    """把文档切成小块"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=CHUNK_SIZE,
        chunk_overlap=CHUNK_OVERLAP,
        length_function=len,
        separators=["\n\n", "\n", "。", "！", "？", "；", "，", " ", ""],
    )
    chunks = splitter.split_documents(documents)
    print(f"  切分: {len(documents)} 段 -> {len(chunks)} 个文本块")
    return chunks


def create_vector_store(chunks, embedding_model):
    """创建并持久化向量数据库"""
    if Path(VECTOR_STORE_DIR).exists():
        shutil.rmtree(VECTOR_STORE_DIR)
        print("  已删除旧的向量数据库")

    print("  正在向量化并存入 ChromaDB...")
    vector_store = Chroma.from_documents(
        documents=chunks,
        embedding=embedding_model,
        persist_directory=VECTOR_STORE_DIR,
    )
    print(f"  完成: 共 {len(chunks)} 个向量")
    return vector_store


def ingest(docs_dir: str = None, vector_dir: str = None):
    """
    主流程：加载 -> 切分 -> 向量化 -> 存储

    参数:
        docs_dir: 文档目录（默认使用 config.DOCUMENTS_DIR）
        vector_dir: 向量数据库目录（默认使用 config.VECTOR_STORE_DIR）
    """
    if docs_dir is None:
        docs_dir = DOCUMENTS_DIR
    if vector_dir is None:
        vector_dir = VECTOR_STORE_DIR

    print("=" * 50)
    print(f"RAG 知识库 - 文档导入工具")
    print(f"  文档目录: {docs_dir}")
    print(f"  向量目录: {vector_dir}")
    print("=" * 50)

    # 1. 加载文档
    print("\n[1/4] 加载文档...")
    documents = load_documents(docs_dir)

    if not documents:
        print(f"\n  {docs_dir} 中没有找到支持的文档！")
        print(f"  请把 PDF/TXT/MD/DOCX 文件放进去")
        return None

    # 2. 切分文档
    print(f"\n[2/4] 文档切分 ({len(documents)} 段)...")
    chunks = split_documents(documents)

    # 3. 加载 Embedding 模型
    print(f"\n[3/4] 加载 Embedding 模型...")
    embedding_model = get_embedding_model()
    print(f"  模型: {LOCAL_EMBEDDING_MODEL}")

    # 4. 创建向量数据库（使用指定目录）
    print(f"\n[4/4] 创建向量数据库...")
    if Path(vector_dir).exists():
        shutil.rmtree(vector_dir)
        print(f"  已删除旧的向量数据库: {vector_dir}")

    print(f"  正在向量化并存入 ChromaDB...")
    vector_store = Chroma.from_documents(
        documents=chunks,
        embedding=embedding_model,
        persist_directory=vector_dir,
    )
    print(f"  完成: 共 {len(chunks)} 个向量")

    # 5. 统计
    source_files = set(d.metadata["source"] for d in documents)
    print("\n" + "=" * 50)
    print("  导入完成！")
    print(f"  文件数: {len(source_files)}")
    print(f"  文本块: {len(chunks)}")
    print(f"  位置: {vector_dir}")
    print("=" * 50)

    return vector_store


if __name__ == "__main__":
    ingest()
