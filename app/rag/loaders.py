"""
文档加载器
每种文件格式有独立的加载逻辑，异常隔离
"""

from __future__ import annotations

from pathlib import Path
from typing import List, Optional

from langchain_core.documents import Document
from langchain_community.document_loaders import PyPDFLoader, TextLoader

from app.logger import logger


def load_pdf(file_path: str) -> List[Document]:
    """加载 PDF 文件"""
    try:
        loader = PyPDFLoader(file_path)
        docs = loader.load()
        logger.debug(f"PDF 加载成功: {Path(file_path).name} -> {len(docs)} 页")
        return docs
    except Exception as e:
        logger.error(f"PDF 加载失败 {file_path}: {e}")
        return []


def load_text(file_path: str) -> List[Document]:
    """加载纯文本文件"""
    try:
        loader = TextLoader(file_path, encoding="utf-8")
        docs = loader.load()
        logger.debug(f"TXT 加载成功: {Path(file_path).name} -> {len(docs)} 段")
        return docs
    except UnicodeDecodeError:
        # 尝试 GBK 编码
        try:
            loader = TextLoader(file_path, encoding="gbk")
            docs = loader.load()
            logger.debug(f"TXT(GBK) 加载成功: {Path(file_path).name}")
            return docs
        except Exception as e:
            logger.error(f"TXT 加载失败 {file_path}: {e}")
            return []
    except Exception as e:
        logger.error(f"TXT 加载失败 {file_path}: {e}")
        return []


def load_markdown(file_path: str) -> List[Document]:
    """加载 Markdown 文件"""
    try:
        from langchain_community.document_loaders import UnstructuredMarkdownLoader
        loader = UnstructuredMarkdownLoader(file_path)
        docs = loader.load()
        logger.debug(f"MD 加载成功: {Path(file_path).name} -> {len(docs)} 段")
        return docs
    except Exception as e:
        logger.error(f"MD 加载失败 {file_path}: {e}")
        return []


def load_docx(file_path: str) -> List[Document]:
    """
    加载 DOCX 文件
    使用 python-docx（比 Unstructured 稳定 10 倍）
    """
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        full_text = "\n".join(paragraphs)
        if not full_text.strip():
            logger.warning(f"DOCX 文件为空: {Path(file_path).name}")
            return []

        logger.debug(f"DOCX 加载成功: {Path(file_path).name} -> {len(full_text)} 字")
        return [Document(
            page_content=full_text,
            metadata={"source": Path(file_path).name},
        )]
    except Exception as e:
        logger.error(f"DOCX 加载失败 {file_path}: {e}")
        return []


# 加载器注册表
LOADER_REGISTRY = {
    ".pdf": load_pdf,
    ".txt": load_text,
    ".md": load_markdown,
    ".docx": load_docx,
}


def load_document(file_path: str) -> List[Document]:
    """
    根据文件扩展名自动选择合适的加载器

    参数:
        file_path: 文件路径

    返回:
        文档块列表，失败返回空列表
    """
    ext = Path(file_path).suffix.lower()
    loader = LOADER_REGISTRY.get(ext)

    if loader is None:
        logger.warning(f"不支持的文件格式: {ext}")
        return []

    docs = loader(file_path)

    # 确保所有文档都有 source 元数据
    filename = Path(file_path).name
    for doc in docs:
        doc.metadata.setdefault("source", filename)

    return docs
